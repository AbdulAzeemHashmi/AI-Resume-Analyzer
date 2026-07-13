import os
import json
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import openai
import PyPDF2
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
import io

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Configure OpenAI
openai.api_key = os.getenv("OPENAI_API_KEY")

# Configure upload folder
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ============================================
# FILE PARSING FUNCTIONS
# ============================================

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        raise Exception(f"Failed to parse PDF: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Failed to parse DOCX: {str(e)}")

def extract_text(file_path, file_type):
    """Extract text based on file type"""
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    else:
        raise Exception(f"Unsupported file type: {file_type}")

# ============================================
# AI FUNCTIONS
# ============================================

def analyze_resume(text):
    """Analyze resume and find weaknesses"""
    try:
        prompt = f"""
        You are a professional career advisor. Review the following resume and identify weaknesses.
        
        Resume:
        {text}
        
        List at least 5 specific weaknesses. Focus on:
        - Weak action verbs
        - Missing quantifiable achievements
        - Poor formatting
        - Irrelevant skills
        - Lack of summary or objective
        
        Return ONLY a JSON array of strings. Example: ["Weakness 1", "Weakness 2", ...]
        """
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional career advisor. Return only JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=500
        )
        
        result = response.choices[0].message.content.strip()
        # Try to parse as JSON
        try:
            weaknesses = json.loads(result)
            if isinstance(weaknesses, list):
                return weaknesses
        except:
            pass
        
        # Fallback: split by lines
        return [line.strip() for line in result.split('\n') if line.strip()]
        
    except Exception as e:
        raise Exception(f"AI analysis failed: {str(e)}")

def enhance_resume(text, weaknesses):
    """Generate improved resume"""
    try:
        weaknesses_text = "\n".join([f"- {w}" for w in weaknesses])
        
        prompt = f"""
        You are a professional resume writer. Improve the following resume.
        
        Original Resume:
        {text}
        
        Weaknesses to Fix:
        {weaknesses_text}
        
        Create an improved version that:
        - Uses strong action verbs
        - Includes quantifiable achievements
        - Has better formatting
        - Removes irrelevant skills
        - Adds a professional summary
        
        Return ONLY the improved resume text. Do not include any additional commentary.
        """
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional resume writer. Return only the improved resume text."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        raise Exception(f"AI enhancement failed: {str(e)}")

# ============================================
# FILE GENERATION FUNCTIONS
# ============================================

def generate_pdf(text):
    """Generate a PDF from text"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Add text
    y = height - inch
    for line in text.split('\n'):
        if y < inch:
            c.showPage()
            y = height - inch
        c.drawString(inch, y, line[:100])  # Limit line length
        y -= 14
    
    c.save()
    buffer.seek(0)
    return buffer

def generate_docx(text):
    """Generate a DOCX from text"""
    doc = docx.Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================
# API ENDPOINTS
# ============================================

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({"status": "healthy"})

@app.route('/upload', methods=['POST'])
def upload_file():
    """Upload and parse a resume file"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Get file extension
        filename = file.filename.lower()
        if filename.endswith('.pdf'):
            file_type = 'pdf'
        elif filename.endswith('.docx'):
            file_type = 'docx'
        else:
            return jsonify({"error": "Unsupported file type. Use PDF or DOCX."}), 400
        
        # Save file temporarily
        temp_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(temp_path)
        
        # Extract text
        text = extract_text(temp_path, file_type)
        
        # Clean up
        os.remove(temp_path)
        
        return jsonify({
            "success": True,
            "text": text,
            "file_type": file_type,
            "filename": file.filename
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/analyze', methods=['POST'])
def analyze():
    """Analyze resume text"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({"error": "Missing text"}), 400
        
        weaknesses = analyze_resume(data['text'])
        return jsonify({
            "success": True,
            "weaknesses": weaknesses
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/enhance', methods=['POST'])
def enhance():
    """Generate improved resume"""
    try:
        data = request.get_json()
        if not data or 'text' not in data or 'weaknesses' not in data:
            return jsonify({"error": "Missing text or weaknesses"}), 400
        
        improved_text = enhance_resume(data['text'], data['weaknesses'])
        return jsonify({
            "success": True,
            "improved_text": improved_text
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/generate', methods=['POST'])
def generate_file():
    """Generate output file"""
    try:
        data = request.get_json()
        if not data or 'text' not in data or 'file_type' not in data:
            return jsonify({"error": "Missing text or file_type"}), 400
        
        text = data['text']
        file_type = data['file_type']
        
        if file_type == 'pdf':
            buffer = generate_pdf(text)
            return send_file(
                buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name='improved_resume.pdf'
            )
        elif file_type == 'docx':
            buffer = generate_docx(text)
            return send_file(
                buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name='improved_resume.docx'
            )
        else:
            return jsonify({"error": "Unsupported file type"}), 400
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)